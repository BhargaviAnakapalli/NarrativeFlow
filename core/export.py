# core/export.py — PDF and DOCX export helpers (renamed from core/export_utils.py)
import io
from fpdf import FPDF
from docx import Document


def _format_paragraphs(text: str) -> list[str]:
    """Split text gracefully into paragraphs to avoid blank massive lines."""
    return [p.strip() for p in text.split("\n") if p.strip()]


def safe_text(text: str) -> str:
    """Strip or replace Unicode characters that FPDF's default latin-1 fonts cannot handle."""
    replacements = {
        '\u201c': '"', '\u201d': '"', '\u2018': "'", '\u2019': "'", '\u2026': "...", '\u2014': "-", '\u2013': "-"
    }
    for search, replace in replacements.items():
        text = text.replace(search, replace)
    return text.encode('latin-1', 'replace').decode('latin-1')


def create_pdf_bytes(title: str, chapter_order: list[str], chapters: dict) -> bytes:
    """Generate a PDF directly into memory and return the bytes."""
    class StoryPDF(FPDF):
        def header(self):
            self.set_font('helvetica', 'I', 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, safe_text(title), border=0, align='R')
            self.ln(15)

        def footer(self):
            self.set_y(-15)
            self.set_font('helvetica', 'I', 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

    pdf = StoryPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("helvetica", "B", 24)
    pdf.cell(0, 20, safe_text(title), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)

    for ch in chapter_order:
        content = chapters.get(ch, "")
        if not content.strip():
            continue

        pdf.set_font("helvetica", "B", 16)
        pdf.cell(0, 10, safe_text(ch), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)

        pdf.set_font("helvetica", "", 12)
        paragraphs = _format_paragraphs(content)
        for p in paragraphs:
            pdf.multi_cell(0, 7, safe_text(p))
            pdf.ln(3)

        pdf.add_page()

    return bytes(pdf.output())


def create_docx_bytes(title: str, chapter_order: list[str], chapters: dict) -> bytes:
    """Generate a DOCX document directly into memory and return the bytes."""
    doc = Document()
    doc.add_heading(title, 0)

    for ch in chapter_order:
        content = chapters.get(ch, "")
        if not content.strip():
            continue

        doc.add_heading(ch, level=1)
        paragraphs = _format_paragraphs(content)
        for p in paragraphs:
            doc.add_paragraph(p)

        doc.add_page_break()

    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()
